"""
CV Generator B2B v3.0
Generator profesjonalnych CV w formacie B2B Network z pełnym wsparciem dla header/footer.
Obsługa języków: Polski (pl) i English (en)
"""

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from lxml import etree
import shutil
import os

# Kolory firmowe B2B Network
COLOR_HEADER = RGBColor(225, 79, 79)  # Czerwony firmowy
COLOR_TEXT = RGBColor(55, 53, 53)     # Ciemnoszary

# Namespace dla XML
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Tłumaczenia
TRANSLATIONS = {
    'pl': {
        'why': 'DLACZEGO',
        'education': 'EDUKACJA',
        'dates': 'Daty',
        'education_header': 'Edukacja',
        'skills': 'UMIEJĘTNOŚCI',
        'certifications': 'CERTYFIKATY',
        'languages': 'JĘZYKI',
        'experience': 'DOŚWIADCZENIE',
        'company_name': 'Nazwa firmy:',
        'position': 'Stanowisko:',
        'responsibilities': 'Zakres obowiązków:',
        'technologies': 'Technologie:',
        'rodo': (
            'Wyrażam zgodę na przetwarzanie moich danych osobowych zawartych w przekazanych '
            'przeze mnie dokumentach przez B2B.net S.A. w celach związanych z moim udziałem '
            'w niniejszym procesie rekrutacyjnym. Ponadto przyjmuję do wiadomości i oświadczam, '
            'że zrozumiałem/am, iż administratorem moich danych osobowych zebranych na podstawie '
            'niniejszej zgody jest B2B.net S.A. z siedzibą w Warszawie, Al. Jerozolimskie 180, '
            '02-486 Warszawa. Dane będą przetwarzane zgodnie z przepisami Rozporządzenia Parlamentu '
            'Europejskiego i Rady (UE) 2016/679 z dnia 27 kwietnia 2016 r. w sprawie ochrony osób '
            'fizycznych w związku z przetwarzaniem danych osobowych i w sprawie swobodnego przepływu '
            'takich danych (dalej „RODO"). Dane przekazałem/am dobrowolnie, przy czym przysługuje mi '
            'prawo do cofnięcia zgody na przetwarzanie danych w dowolnym momencie poprzez wysłanie '
            'żądania na adres: rekrutacja@b2bnetwork.pl. Podanie danych jest niezbędne do realizacji '
            'ww. celu, dlatego żądanie ich usunięcia jest równoznaczne z rezygnacją z dalszego udziału '
            'w procesie rekrutacyjnym. Przysługuje mi również prawo dostępu do treści moich danych oraz '
            'ich poprawiania w każdym czasie. Jestem również świadomy/świadoma, że odbiorcami moich danych '
            'osobowych mogą być wyłącznie podmioty upoważnione na podstawie przepisów prawa, a także '
            'upoważnione na podstawie umów zawartych przez B2B.net S.A., w szczególności z klientami.'
        )
    },
    'en': {
        'why': 'WHY',
        'education': 'EDUCATION',
        'dates': 'Dates',
        'education_header': 'Education',
        'skills': 'SKILLS',
        'certifications': 'CERTIFICATIONS',
        'languages': 'LANGUAGES',
        'experience': 'EXPERIENCE',
        'company_name': 'Company:',
        'position': 'Position:',
        'responsibilities': 'Responsibilities:',
        'technologies': 'Technologies:',
        'rodo': (
            'I hereby consent to the processing of my personal data contained in the documents '
            'submitted by me by B2B.net S.A. for purposes related to my participation in this '
            'recruitment process. Furthermore, I acknowledge and declare that I have understood '
            'that the administrator of my personal data collected on the basis of this consent is '
            'B2B.net S.A. with its registered office in Warsaw, Al. Jerozolimskie 180, 02-486 Warsaw. '
            'The data will be processed in accordance with the provisions of Regulation (EU) 2016/679 '
            'of the European Parliament and of the Council of 27 April 2016 on the protection of natural '
            'persons with regard to the processing of personal data and on the free movement of such data '
            '(hereinafter "GDPR"). I have provided the data voluntarily, and I have the right to withdraw '
            'my consent to data processing at any time by sending a request to: rekrutacja@b2bnetwork.pl. '
            'The provision of data is necessary for the realization of the above purpose, therefore requesting '
            'their deletion is tantamount to resignation from further participation in the recruitment process. '
            'I also have the right to access the content of my data and to correct them at any time. I am also '
            'aware that the recipients of my personal data may only be entities authorized under the law, as well '
            'as authorized under contracts concluded by B2B.net S.A., in particular with customers.'
        )
    }
}


def add_horizontal_line(doc):
    """Dodaje czerwoną linię poziomą - element brandingu B2B"""
    para = doc.add_paragraph()
    
    hr_xml = '''
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:v="urn:schemas-microsoft-com:vml"
         xmlns:o="urn:schemas-microsoft-com:office:office">
        <w:rPr><w:noProof/></w:rPr>
        <w:pict>
            <v:rect style="width:448.6pt;height:2pt" fillcolor="#e14f4f" stroked="f" 
                    o:hr="t" o:hrstd="t" o:hrnoshade="t" o:hrpct="989" o:hralign="center"/>
        </w:pict>
    </w:r>
    '''
    
    hr_element = parse_xml(hr_xml)
    para._element.append(hr_element)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_section_header(doc, text):
    """Dodaje nagłówek sekcji z czerwoną linią"""
    add_horizontal_line(doc)

    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Montserrat SemiBold'
    run.font.color.rgb = COLOR_HEADER
    run.font.bold = False
    run.font.size = Pt(14)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(3)
    return para


def add_bullet_point(doc, text, punctuation=''):
    """Dodaje punkt wypunktowany z WŁAŚCIWĄ NUMERACJĄ WORD."""
    # Usuń istniejącą interpunkcję na końcu i dodaj nową
    text = text.rstrip('.,;')
    text = text + punctuation

    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Montserrat'
    run.font.color.rgb = COLOR_TEXT
    run.font.size = Pt(10)

    pPr = para._element.get_or_add_pPr()
    numPr = etree.Element(f'{{{NS_W}}}numPr')

    ilvl = etree.SubElement(numPr, f'{{{NS_W}}}ilvl')
    ilvl.set(f'{{{NS_W}}}val', '0')

    numId = etree.SubElement(numPr, f'{{{NS_W}}}numId')
    numId.set(f'{{{NS_W}}}val', '1')

    pPr.insert(0, numPr)

    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing = 1.5
    return para


def add_bullet_list(doc, items):
    """Dodaje listę punktowaną z przecinkami i kropką na końcu."""
    for i, item in enumerate(items):
        is_last = (i == len(items) - 1)
        punctuation = '.' if is_last else ','
        add_bullet_point(doc, item, punctuation)


def create_cv_from_template(candidate_data, template_path, output_path):
    """
    Tworzy CV rozpoczynając OD SZABLONU.

    Args:
        candidate_data: Dict z danymi kandydata (zawiera pole 'language': 'pl' lub 'en')
        template_path: Ścieżka do szablon_firmowy.docx
        output_path: Gdzie zapisać wynikowe CV
    """

    # Pobierz język z danych kandydata (domyślnie 'pl')
    language = candidate_data.get('language', 'pl')
    blind_cv = candidate_data.get('blind_cv', False)
    t = TRANSLATIONS.get(language, TRANSLATIONS['pl'])

    # Anonimizacja dla Blind CV
    if blind_cv:
        if language == 'en':
            candidate_data['name'] = 'Candidate'
            candidate_data['first_name'] = 'Candidate'
        else:
            candidate_data['name'] = 'Kandydat'
            candidate_data['first_name'] = 'Kandydat'

        # Anonimizuj nazwy firm w doświadczeniu
        for job in candidate_data.get('experience', []):
            industry = job.get('industry', 'IT')
            if language == 'en':
                job['company'] = f'Company from {industry} industry'
            else:
                job['company'] = f'Firma z branży {industry}'

    print(f"Generuję CV dla {candidate_data['name']} (język: {language}, blind: {blind_cv})...")
    
    # KROK 1: Skopiuj szablon jako bazę
    temp_cv = '/tmp/cv_base.docx'
    shutil.copy(template_path, temp_cv)
    print("  ✓ Skopiowano szablon jako bazę")
    
    # KROK 2: Otwórz skopiowany szablon
    doc = Document(temp_cv)
    
    # KROK 3: Wyczyść treść (ale ZACHOWAJ sectPr!)
    for element in list(doc.element.body):
        if not element.tag.endswith('}sectPr'):
            doc.element.body.remove(element)
    print("  ✓ Wyczyszczono zawartość szablonu")
    
    # KROK 4: Dodaj treść CV
    
    # === NAGŁÓWEK GŁÓWNY ===
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(
        f"{candidate_data['position']} – {candidate_data['name']}"
    )
    header_run.font.name = 'Montserrat SemiBold'
    header_run.font.color.rgb = COLOR_HEADER
    header_run.font.bold = False
    header_run.font.size = Pt(24)
    header_para.paragraph_format.space_after = Pt(6)
    header_para.paragraph_format.space_before = Pt(2)
    
    add_horizontal_line(doc)
    
    # === DLACZEGO [IMIĘ] / WHY [NAME] ===
    para = doc.add_paragraph()
    run = para.add_run(f"{t['why']} {candidate_data['first_name'].upper()}")
    run.font.name = 'Montserrat SemiBold'
    run.font.color.rgb = COLOR_HEADER
    run.font.bold = False
    run.font.size = Pt(14)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)
    
    add_bullet_list(doc, candidate_data.get('why_points', []))
    
    # === EDUKACJA / EDUCATION ===
    add_section_header(doc, t['education'])
    
    table = doc.add_table(rows=len(candidate_data['education']) + 1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(5.0)

    # Ustaw właściwości tabeli - usunięcie zewnętrznych obramowań
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # Obramowanie tabeli - wszystkie krawędzie
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Cell margins dla lepszego odstępu
    tblCellMar = OxmlElement('w:tblCellMar')
    for margin_name in ['top', 'bottom']:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), '80')
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)
    for margin_name in ['left', 'right']:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), '120')
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)
    tblPr.append(tblCellMar)

    header_cells = table.rows[0].cells
    header_cells[0].text = t['dates']
    header_cells[1].text = t['education_header']

    for cell in header_cells:
        # Szare tło nagłówka
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E8E8E8')
        shading_elm.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading_elm)

        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                run.font.name = 'Montserrat'
                run.font.bold = True
                run.font.color.rgb = COLOR_TEXT
                run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, edu in enumerate(candidate_data['education'], start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = edu['dates']
        row_cells[1].text = f"{edu['institution']}\n{edu['degree']}\n{edu['location']}"

        # Alternatywne kolorowanie wierszy (zebra striping) - delikatny szary
        if i % 2 == 0:
            for cell in row_cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F8F8F8')
                shading_elm.set(qn('w:val'), 'clear')
                cell._element.get_or_add_tcPr().append(shading_elm)

        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(6)
                for run in paragraph.runs:
                    run.font.name = 'Montserrat'
                    run.font.color.rgb = COLOR_TEXT
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # === UMIEJĘTNOŚCI / SKILLS ===
    add_section_header(doc, t['skills'])
    
    for skill_category in candidate_data.get('skills', []):
        para = doc.add_paragraph()
        
        pPr = para._element.get_or_add_pPr()
        numPr = etree.Element(f'{{{NS_W}}}numPr')
        ilvl = etree.SubElement(numPr, f'{{{NS_W}}}ilvl')
        ilvl.set(f'{{{NS_W}}}val', '0')
        numId = etree.SubElement(numPr, f'{{{NS_W}}}numId')
        numId.set(f'{{{NS_W}}}val', '1')
        pPr.insert(0, numPr)
        
        run1 = para.add_run(skill_category['label'] + ' ')
        run1.font.name = 'Montserrat'
        run1.font.bold = True
        run1.font.color.rgb = COLOR_TEXT
        run1.font.size = Pt(10)
        
        run2 = para.add_run(skill_category['content'])
        run2.font.name = 'Montserrat'
        run2.font.color.rgb = COLOR_TEXT
        run2.font.size = Pt(10)
        
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
    
    # === CERTYFIKATY / CERTIFICATIONS ===
    add_section_header(doc, t['certifications'])

    add_bullet_list(doc, candidate_data.get('certifications', []))

    # === JĘZYKI / LANGUAGES ===
    add_section_header(doc, t['languages'])

    add_bullet_list(doc, candidate_data.get('languages', []))
    
    # === DOŚWIADCZENIE / EXPERIENCE ===
    add_section_header(doc, t['experience'])
    
    for i, job in enumerate(candidate_data.get('experience', [])):
        if i > 0:
            add_horizontal_line(doc)
        
        # Data
        para = doc.add_paragraph()
        run = para.add_run(job['dates'])
        run.font.name = 'Montserrat'
        run.font.color.rgb = COLOR_TEXT
        run.font.size = Pt(10)
        run.font.bold = True
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(0)
        
        # Nazwa firmy
        para = doc.add_paragraph()
        run1 = para.add_run(t['company_name'] + ' ')
        run1.font.name = 'Montserrat'
        run1.font.color.rgb = COLOR_TEXT
        run1.font.size = Pt(10)
        run2 = para.add_run(job['company'])
        run2.font.name = 'Montserrat'
        run2.font.color.rgb = COLOR_TEXT
        run2.font.size = Pt(10)
        run2.font.bold = True
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        # Stanowisko
        para = doc.add_paragraph()
        run1 = para.add_run(t['position'] + ' ')
        run1.font.name = 'Montserrat'
        run1.font.color.rgb = COLOR_TEXT
        run1.font.size = Pt(10)
        run2 = para.add_run(job['position'])
        run2.font.name = 'Montserrat'
        run2.font.color.rgb = COLOR_TEXT
        run2.font.size = Pt(10)
        run2.font.bold = True
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        
        # Zakres obowiązków
        para = doc.add_paragraph()
        run = para.add_run(t['responsibilities'])
        run.font.name = 'Montserrat'
        run.font.color.rgb = COLOR_TEXT
        run.font.size = Pt(10)
        run.font.bold = True
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(1)
        
        add_bullet_list(doc, job.get('responsibilities', []))
        
        # Technologie
        if job.get('technologies'):
            para = doc.add_paragraph()
            run1 = para.add_run(t['technologies'] + ' ')
            run1.font.name = 'Montserrat'
            run1.font.color.rgb = COLOR_TEXT
            run1.font.size = Pt(10)
            run1.font.bold = True
            run2 = para.add_run(', '.join(job['technologies']))
            run2.font.name = 'Montserrat'
            run2.font.color.rgb = COLOR_TEXT
            run2.font.size = Pt(10)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(2)
    
    # === KLAUZULA RODO / GDPR ===
    doc.add_paragraph()
    doc.add_paragraph()
    
    rodo_para = doc.add_paragraph()
    rodo_text = t['rodo']
    
    run = rodo_para.add_run(rodo_text)
    run.font.name = 'Montserrat'
    run.font.size = Pt(5)
    run.font.color.rgb = COLOR_TEXT
    
    # KROK 5: Zapisz dokument
    doc.save(output_path)
    
    # Cleanup
    if os.path.exists(temp_cv):
        os.remove(temp_cv)
    
    print(f"  ✓ Zapisano CV: {output_path}")
    print("\n✅ CV wygenerowane pomyślnie!")
    print(f"   ✓ Język: {language.upper()}")
    print("   ✓ Header z logo B2B Network")
    print("   ✓ Footer z logo B2B Network")
    print("   ✓ Numeracja Word dla bullet points")
    print("   ✓ Wszystkie sekcje obecne")
    
    return doc


# CLI Interface
if __name__ == '__main__':
    import sys
    import json
    
    if len(sys.argv) != 4:
        print("Usage: python generate_cv.py <data_json_path> <template_path> <output_path>")
        sys.exit(1)
    
    data_json_path = sys.argv[1]
    template_path = sys.argv[2]
    output_path = sys.argv[3]
    
    with open(data_json_path, 'r', encoding='utf-8') as f:
        candidate_data = json.load(f)
    
    create_cv_from_template(candidate_data, template_path, output_path)
